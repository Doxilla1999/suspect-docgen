import os
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
GARUDA_PATH = os.path.join(SCRIPT_DIR, "garuda.png")
PLACEHOLDER_DIR = os.path.join(SCRIPT_DIR, "placeholders")

FONT = "TH SarabunIT๙"
SIZE = Pt(16)

def set_font(run, bold=False, size=SIZE):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:cs'), FONT)

def add_p(doc, text="", bold=False, align=None, size=SIZE, space_after=2):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.0
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold, size=size)
    return p

def add_mixed(doc, parts, align=None, space_after=2):
    # parts: list of (text, bold)
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.0
    for text, bold in parts:
        r = p.add_run(text)
        set_font(r, bold=bold)
    return p

def base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.2)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = SIZE
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(1)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    return doc

def letterhead(doc, doc_number_tag, station_tag="{station_name}", province_tag="{province}", postal_tag="{postal_code}", date_tag="{doc_date}"):
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(0)
    run0 = p0.add_run()
    run0.add_picture(GARUDA_PATH, height=Cm(3.0))
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.tab_stops.add_tab_stop(Cm(17.3), WD_TAB_ALIGNMENT.RIGHT)
    r_label = p1.add_run("ที่ ")
    set_font(r_label, bold=True)
    r_num = p1.add_run(doc_number_tag)
    set_font(r_num)
    r_tab = p1.add_run("\t")
    set_font(r_tab)
    r_station = p1.add_run(station_tag)
    set_font(r_station)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"จังหวัด{province_tag}  {postal_tag}")
    set_font(r2)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(date_tag)
    set_font(r3)
    doc.add_paragraph()

def field_line(doc, label, tag, trailing=""):
    add_mixed(doc, [(label, True), (" " + tag + trailing, False)])

def signature_block(doc, rank_tag, name_tag, position_tag):
    add_p(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(rank_tag)
    set_font(r)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("( " + name_tag + " )")
    set_font(r2)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(position_tag)
    set_font(r3)

OUT = os.path.join(SCRIPT_DIR, "build") + os.sep
os.makedirs(OUT, exist_ok=True)

# ---------- 1. urine referral ----------
doc = base_doc()
letterhead(doc, "{doc_number}")
add_mixed(doc, [("เรื่อง\t", True), ("ขอส่งตรวจยืนยัน/สารเสพติดในปัสสาวะ ในขั้นที่สองด้วยหลักการทางวิทยาศาสตร์", False)])
add_mixed(doc, [("เรียน\t", True), ("ผู้อำนวยการโรงพยาบาล{hospital_name}", False)])
add_p(doc, "เนื่องด้วย {station_name} ขอส่งตรวจยืนยัน{drug_type}/สารเสพติดในปัสสาวะ ในขั้นที่สองด้วยหลักการทางวิทยาศาสตร์ "
          "เพื่อเป็นการดำเนินการตามแนวทางการตรวจพิสูจน์หาสารเสพติดในปัสสาวะตามพระราชบัญญัติฟื้นฟูสมรรถภาพผู้ติดยาเสพติด พ.ศ. ๒๕๔๕ "
          "ที่ห้องเคมีคลินิกและพิษวิทยา กลุ่มงานพยาธิวิทยาคลินิก โรงพยาบาล{hospital_name} ในวันที่ {test_date} เป็นจำนวน {test_count} ราย ดังมีรายชื่อต่อไปนี้")
add_p(doc, "{suspect_list_block}")
add_p(doc, "ทั้งนี้หากผลการตรวจยืนยันการคัดกรองเป็นอย่างไรขอให้ดำเนินการแจ้งให้ทราบ")
add_p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา")
add_p(doc, "ขอแสดงความนับถือ")
signature_block(doc, "{signer_rank}", "{signer_name}", "{signer_position}")
add_p(doc, "")
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
rows_data = [
    ("สำหรับการรับ-ส่งตัวอย่าง", "สำหรับรับผลการตรวจวิเคราะห์"),
    ("ผู้ส่ง", "ผู้มอบผล"),
    ("ผู้รับ", "ผู้รับผล"),
    ("วันที่", "วันที่"),
    ("เวลา         น.", "เวลา         น."),
]
for i, (a, b) in enumerate(rows_data):
    ra, rb = table.rows[i].cells[0].paragraphs[0].add_run(a), table.rows[i].cells[1].paragraphs[0].add_run(b)
    set_font(ra, bold=(i == 0))
    set_font(rb, bold=(i == 0))
doc.save(OUT + "urine_referral_template.docx")

# ---------- 2. drug test record (pys115) ----------
CB_ON = "☑"
CB_OFF = "☐"
def cb(flag_tag):
    return "{%s}" % flag_tag  # resolved client-side to CB_ON/CB_OFF

doc = base_doc()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("แบบ ปยส. 115"); set_font(r, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("แบบบันทึกการตรวจหรือทดสอบสารเสพติดในร่างกาย"); set_font(r, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("การตรวจสอบพฤติการณ์และสอบถามความสมัครใจเข้ารับการบำบัดรักษา"); set_font(r, bold=True)
field_line(doc, "สถานที่ตรวจ", "{station_name}")
field_line(doc, "วันที่", "{record_day} เดือน {record_month} พ.ศ. {record_year}")
add_p(doc, "ข้าพเจ้า {officer1_rank}{officer1_name} ตำแหน่ง {officer1_position}")
field_line(doc, "สังกัด", "{officer_affiliation}")
p_card1 = add_p(doc, f"{cb('cb_card_ppst')} บัตรประจำตัวเจ้าพนักงาน ป.ป.ส.\tเลขที่ {{card_no_ppst}} หรือ")
p_card1.paragraph_format.tab_stops.add_tab_stop(Cm(10.3), WD_TAB_ALIGNMENT.LEFT)
p_card2 = add_p(doc, f"{cb('cb_card_gov')} บัตรประจำตัวเจ้าหน้าที่ของรัฐ\tเลขที่ {{card_no_gov}}")
p_card2.paragraph_format.tab_stops.add_tab_stop(Cm(10.3), WD_TAB_ALIGNMENT.LEFT)
add_p(doc, "โดยอาศัยอำนาจตามมาตรา 115 แห่งประมวลกฎหมายยาเสพติด")
add_p(doc, "ได้สอบถาม {suspect_title}{suspect_full_name} อายุ {suspect_age} ปี")
add_p(doc, f"{cb('cb_id_card')} บัตรประชาชน  {cb('cb_id_alien')} บัตรคนซึ่งไม่มีสัญชาติไทย  {cb('cb_id_passport')} หนังสือเดินทาง  {cb('cb_id_other')} เอกสารอื่นที่ราชการออกให้ ระบุ {{id_other_detail}}")
field_line(doc, "เลขที่", "{suspect_id_number}")
field_line(doc, "ที่อยู่", "{suspect_address}")
field_line(doc, "ที่อยู่ปัจจุบัน", "{suspect_current_address}  หมายเลขโทรศัพท์ {suspect_phone}")
p = doc.add_paragraph(); r = p.add_run("ส่วนที่ 1 การตรวจหรือทดสอบสารเสพติดในร่างกาย"); set_font(r, bold=True)
add_p(doc, f"ผลการตรวจ  {cb('cb_test_pos')} ผลบวก หมายถึง ทดสอบเบื้องต้นพบว่าอาจมีสารเสพติดอยู่ในร่างกาย คือ {{drug_type}}")
add_p(doc, f"{cb('cb_test_neg')} ผลลบ หมายถึง ทดสอบเบื้องต้นไม่พบว่ามีสารเสพติดอยู่ในร่างกาย")
p = doc.add_paragraph(); r = p.add_run("ผลการตรวจหรือค้น"); set_font(r, bold=True)
add_p(doc, f"{cb('cb_search_none')} ไม่พบยาเสพติด")
add_p(doc, f"{cb('cb_search_found')} พบและยึดยาเสพติดประเภท/ชนิด {{search_drug_type}} ปริมาณ {{search_amount}}")
add_p(doc, "โดยได้ทำบันทึกการตรวจยึดเพื่อดำเนินการตามกฎหมายต่อไปแล้ว")
p = doc.add_paragraph(); r = p.add_run("ส่วนที่ 2 การตรวจสอบพฤติการณ์"); set_font(r, bold=True)
p = doc.add_paragraph(); r = p.add_run("2.1 การตรวจสอบพฤติการณ์ต้องห้ามตามประมวลกฎหมายยาเสพติด มาตรา 115"); set_font(r, bold=True)
add_p(doc, f"{cb('cb_pb1')} อยู่ระหว่างตกเป็นผู้ต้องหาหรืออยู่ในระหว่างถูกดำเนินคดีในความผิดอื่นซึ่งเป็นความผิดที่มีโทษจำคุก")
add_p(doc, f"{cb('cb_pb2')} อยู่ในระหว่างรับโทษจำคุกตามคำพิพากษาของศาล")
add_p(doc, f"{cb('cb_pb3')} มีพฤติกรรมที่อาจก่อให้เกิดอันตรายแก่ผู้อื่นหรือสังคม หรือมีพฤติกรรมที่อาจก่อให้เกิดอันตรายแก่ผู้อื่นหรือสังคมที่เกิดจากโรคทางจิตและประสาท หรืออาการที่เกิดจากฤทธิ์ของยาเสพติดที่ใช้ เช่น สมัครใจเข้ารับการบำบัดรักษา แต่ไม่ให้ความร่วมมือหรือหลบหนี หรือมีลักษณะเป็นอันธพาลที่ก่อให้เกิดความรุนแรงกับผู้อื่นหรือสังคม")
add_p(doc, f"{cb('cb_pb4')} ไม่พบพฤติการณ์ต้องห้ามตามกฎหมาย")
p = doc.add_paragraph(); r = p.add_run("2.2 การตรวจสอบข้อมูลส่วนบุคคลเกี่ยวกับอาชีพรายได้"); set_font(r, bold=True)
field_line(doc, "อาชีพ", "{suspect_occupation}  รายได้โดยประมาณ {suspect_income}")
p = doc.add_paragraph(); r = p.add_run("2.3 พฤติการณ์อื่นที่เกี่ยวข้องกับยาเสพติด (ถ้ามี)"); set_font(r, bold=True)
add_p(doc, "{other_notes}")
p = doc.add_paragraph(); r = p.add_run("ส่วนที่ 3 การสอบถามความสมัครใจเข้ารับการบำบัดรักษา"); set_font(r, bold=True)
add_p(doc, "ข้าพเจ้า (ชื่อสกุล) {suspect_full_name}")
add_p(doc, f"{cb('cb_consent_yes')} ขอลงนามสมัครใจเข้ารับการบำบัดรักษา")
add_p(doc, f"{cb('cb_consent_no')} ขอลงนามไม่สมัครใจเข้ารับการบำบัดรักษา")
add_p(doc, size=Pt(13), text="กรณีสมัครใจเข้ารับการบำบัดรักษาจะต้องปฏิบัติครบถ้วนตามหลักเกณฑ์ วิธีการ และเงื่อนไขที่คณะกรรมการบำบัดรักษาและฟื้นฟูผู้ติดยาเสพติดกำหนด "
          "จนได้รับการรับรองเป็นหนังสือว่าเป็นผู้ผ่านการบำบัดรักษาเป็นที่น่าพอใจจากหัวหน้าสถานพยาบาลยาเสพติดหรือสถานฟื้นฟูสมรรถภาพผู้ติดยาเสพติด")
add_p(doc, size=Pt(13), text="ทั้งนี้ ในกรณีที่ไม่ไปเข้ารับการบำบัดรักษา หลบหนี หรือไม่ได้รับการรับรองเป็นหนังสือว่าเป็นผู้ผ่านการบำบัดรักษาเป็นที่น่าพอใจ "
          "หากเจ้าพนักงานตรวจพบว่าเสพหรือครอบครองเพื่อเสพอีกจะไม่สามารถสมัครใจเข้ารับการบำบัดรักษาในครั้งต่อไปได้ และจะต้องถูกดำเนินคดีตามกฎหมาย "
          "ซึ่งมีอัตราโทษจำคุกสูงสุดไม่เกิน ๒ ปี หรือปรับไม่เกิน ๔๐,๐๐๐ บาท หรือทั้งจำทั้งปรับ แล้วแต่ฐานความผิดตามประมวลกฎหมายยาเสพติด")
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.0
p = doc.add_paragraph(); r = p.add_run("ส่วนที่ 4 การนัดหมายกรณีไม่สามารถส่งตัวไปคัดกรองยังสถานพยาบาลยาเสพติด/ศูนย์คัดกรอง (ถ้ามี)"); set_font(r, bold=True)
add_p(doc, "เจ้าพนักงานนัดหมายให้เข้ารับการคัดกรองในวันที่ {appointment_date} เวลา {appointment_time} น. ณ สถานที่ {appointment_place}")
add_p(doc, size=Pt(13), text="กรณีไม่มารายงานตัวและรับการคัดกรองในวันและเวลาที่เจ้าพนักงานกำหนด จะถือว่าไม่สมัครใจเข้ารับการบำบัดรักษาและจะต้องถูกดำเนินคดีตามกฎหมาย "
          "ซึ่งมีอัตราโทษจำคุกสูงสุดไม่เกิน ๒ ปี หรือปรับไม่เกิน ๔๐,๐๐๐ บาท หรือทั้งจำทั้งปรับ แล้วแต่ฐานความผิดตามประมวลกฎหมายยาเสพติด")
p = doc.add_paragraph(); r = p.add_run("ส่วนที่ 5 การรับรองและให้ความยินยอม"); set_font(r, bold=True)
add_p(doc, size=Pt(13), text="ข้าพเจ้าขอรับรองว่าผลการตรวจหรือทดสอบสารเสพติดในปัสสาวะเบื้องต้นและผลการตรวจค้น/ยึดยาเสพติดตามที่ปรากฏ "
          "เจ้าหน้าที่ได้ตรวจสอบต่อหน้าข้าพเจ้าและเป็นความจริงทุกประการ")
add_p(doc, size=Pt(13), text="ในระหว่างที่เจ้าพนักงานให้ข้าพเจ้าอยู่ในความดูแลเป็นการชั่วคราว ได้จัดสถานที่ราชการหรือสถานที่อื่นใดที่เจ้าพนักงานเห็นสมควร "
          "โดยไม่ปะปนผู้ต้องหาหรือมีบุคคลอื่นที่ไม่เกี่ยวข้องอยู่ในสถานที่นั้น อันมีลักษณะเป็นการประจาน และคำนึงถึงอายุ เพศ และสภาวะเป็นสำคัญแล้ว")
add_p(doc, size=Pt(13), text="ข้าพเจ้ายินยอมให้หน่วยงานรัฐบันทึก/ใช้/แลกเปลี่ยน/เปิดเผยข้อมูลส่วนบุคคลของข้าพเจ้า เพื่อประโยชน์ในการพิจารณาส่งตัวเข้าสู่กระบวนการบำบัดรักษา "
          "รวมถึงจัดทำนโยบายและแผนว่าด้วยการป้องกัน ปราบปราม และแก้ไขปัญหายาเสพติด ทั้งนี้การถ่ายสำเนา ถ่ายภาพ หรือบันทึกไว้ไม่ว่าในรูปแบบใดๆ "
          "ให้ถือเป็นหลักฐานในการให้ความยินยอมของข้าพเจ้าเช่นเดียวกัน")
add_p(doc, "ข้าพเจ้าได้อ่าน/เจ้าพนักงานได้อ่านบันทึกนี้ให้ฟังแล้ว รับรองว่าถูกต้องจึงลงลายมือชื่อไว้เป็นหลักฐาน")
add_p(doc, "")
def sig_line(text, name_tag):
    add_p(doc, "(ลงชื่อ) ............................................ " + text)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("( " + name_tag + " )"); set_font(r)
sig_line("ผู้เข้ารับการตรวจค้น/ตรวจหรือทดสอบฯ/ผู้ยินยอมฯ", "{suspect_full_name}")
sig_line("{officer1_rank} เจ้าพนักงาน ป.ป.ส. ผู้ตรวจหรือทดสอบฯ", "{officer1_name}")
sig_line("ผู้สมัครใจ/ไม่สมัครใจเข้ารับการบำบัดรักษา", "{suspect_full_name}")
sig_line("ผู้ช่วยเจ้าพนักงาน ป.ป.ส. (ถ้ามี)", "{officer2_rank}{officer2_name}")
sig_line("พยาน", "{officer3_rank}{officer3_name}")
doc.save(OUT + "drug_test_115_template.docx")

# ---------- 6. combined person profile (page 1) + 4-angle photos (page 2) ----------
from docx.shared import Cm as CmSize
from docx.oxml.ns import qn as _qn

def set_cell_text(tcell, text, bold=False, size=None):
    tcell.text = ""
    p = tcell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_font(r, bold=bold, size=size or SIZE)
    return p

def shade_cell(tcell, hex_color):
    tcPr = tcell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(_qn('w:shd'), {_qn('w:val'): 'clear', _qn('w:color'): 'auto', _qn('w:fill'): hex_color})
    tcPr.append(shd)

doc = base_doc()
doc.sections[0].top_margin = Cm(0.4)
doc.sections[0].bottom_margin = Cm(0.35)
doc.sections[0].left_margin = Cm(1.4)
doc.sections[0].right_margin = Cm(1.2)

add_mixed(doc, [("ปจว. ข้อ", False), ("..................... ", False), ("เวลา", False), ("..................... ", False),
                ("คดีที่", False), ("..................... ", False), ("ลง", False), ("..................... ", False),
                ("พงส.", False), ("..................... ", False)])
add_p(doc, "ประวัติบุคคลที่เกี่ยวข้องกับอาชญากรรม", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_p(doc, "ผนวก ซ.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_p(doc, "{station_name} อำเภอ{station_district} จังหวัด{province}", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_p(doc, 'ประเภทบุคคล : "{charge}"', size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_mixed(doc, [("วัน เดือน ปี ที่จัดทำประวัติ  ", True), ("{record_date}", False), ("      ผู้จัดทำ  ", True), ("{compiler_name}", False)])
add_p(doc, "ประวัติ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

profile_rows = [
    [("๑) ชื่อ", "{suspect_first_name}"), ("นามสกุล", "{suspect_last_name}"), ("ชื่อเล่น", "{suspect_nickname}"), ("นามแฝง", "{suspect_alias}")],
    [("อาชีพ", "{suspect_occupation}"), ("เชื้อชาติ", "{suspect_ethnicity}"), ("สัญชาติ", "{suspect_nationality}"), ("ศาสนา", "{suspect_religion}")],
    [("๒) อายุ", "{suspect_age} ปี"), ("วัน/เดือน/ปี (เกิด)", "{suspect_dob}"), ("", ""), ("", "")],
]
profile_rows2 = [
    [("ตำบล", "{addr_tambon}"), ("อำเภอ", "{addr_amphoe}"), ("จังหวัด", "{addr_province}"), ("", "")],
    [("๔) การศึกษา", "{suspect_education}"), ("", ""), ("", ""), ("", "")],
    [("๕) บิดา มารดา", "{suspect_parents}"), ("", ""), ("", ""), ("", "")],
    [("๖) สามี/ภรรยา", "{suspect_spouse}"), ("ที่อยู่", "{suspect_spouse_address}"), ("", ""), ("", "")],
    [("เบอร์โทร (ผู้ต้องหา)", "{suspect_phone}"), ("", ""), ("", ""), ("", "")],
    [("๗) เพื่อนสนิท หรือผู้ให้พักอาศัย", "{suspect_friends}"), ("ตำหนิรูปพรรณ", ""), ("", ""), ("", "")],
    [("๘) ที่อยู่ปัจจุบัน", "{suspect_current_address}"), ("ความสูง", "{phys_height}"), ("", ""), ("", "")],
    [("๙) สถานที่ชอบไปเที่ยวเตร่", "{suspect_hangouts}"), ("น้ำหนัก", "{phys_weight}"), ("", ""), ("", "")],
    [("๑๐) อาวุธที่ใช้ในการกระทำผิด", "{crime_weapon}"), ("สีผิว", "{phys_skin}"), ("", ""), ("", "")],
    [("๑๑) ยานพาหนะที่ใช้", "{crime_vehicle}"), ("ศีรษะ", "{phys_head}"), ("", ""), ("", "")],
    [("๑๒) ผู้ร่วมกระทำผิด", "{crime_accomplice}"), ("ผม", "{phys_hair}"), ("", ""), ("", "")],
]
profile_rows3 = [
    [("ตำรวจหรือบุคคลผู้รู้จักตัว", "{known_by}"), ("จมูก", "{phys_nose}"), ("", ""), ("", "")],
    [("บัตรประจำตัวเลขที่", "{suspect_id_number}"), ("ใบหู", "{phys_ears}"), ("", ""), ("", "")],
    [("ออกให้อำเภอ", "{id_issue_amphoe}"), ("หนวดเครา", "{phys_beard}"), ("", ""), ("", "")],
    [("จังหวัด", "{id_issue_province}"), ("ตำหนิ", "{phys_marks}"), ("", ""), ("", "")],
]

TABS4 = [Cm(2.6), Cm(6.6), Cm(8.6), Cm(12.6), Cm(14.6)]
TABS2 = [Cm(10.3), Cm(12.3)]
PSIZE = Pt(16)

def tabline(doc, segments, tabs, size=PSIZE, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.0
    for pos in tabs:
        p.paragraph_format.tab_stops.add_tab_stop(pos, WD_TAB_ALIGNMENT.LEFT)
    text = "\t".join(segments)
    r = p.add_run(text)
    set_font(r, size=size)
    return p

def field_pair(doc, left_label, left_tag, right_label="", right_tag=""):
    right = f"{right_label}\t{right_tag}" if right_label else ""
    tabline(doc, [f"{left_label} {left_tag}", right] if right else [f"{left_label} {left_tag}"],
            [Cm(10.3), Cm(12.3)])

tabline(doc, ["๑) ชื่อ {suspect_first_name}", "นามสกุล {suspect_last_name}", "นามแฝง {suspect_alias}{suspect_nickname}", "อาชีพ {suspect_occupation}"], TABS4)
tabline(doc, ["เชื้อชาติ {suspect_ethnicity}", "สัญชาติ {suspect_nationality}", "ศาสนา {suspect_religion}", "", ""], TABS4)
tabline(doc, ["๒) อายุ {suspect_age} ปี", "วัน/เดือน/ปี (เกิด) {suspect_dob}", "", "", ""], TABS4)
tabline(doc, ["๓) ภูมิลำเนาครั้งสุดท้ายอยู่บ้านเลขที่ {addr_house_no}", "หมู่ที่ {addr_moo}"], TABS4)
tabline(doc, ["ถนน/ซอย {addr_road}"], TABS4)
tabline(doc, ["ตำบล {addr_tambon}", "อำเภอ {addr_amphoe}", "จังหวัด {addr_province}", "", ""], TABS4)
tabline(doc, ["๔) การศึกษา {suspect_education}"], TABS4)
tabline(doc, ["๕) บิดา มารดา {suspect_parents}"], TABS4)
tabline(doc, ["๖) สามี/ภรรยา ชื่อ {suspect_spouse}"], TABS4)
tabline(doc, ["ที่อยู่ {suspect_spouse_address}"], TABS4)
tabline(doc, [""], TABS4, space_after=1)
tabline(doc, ["เบอร์โทร (ผู้ต้องหา) {suspect_phone}", "", "ตำหนิรูปพรรณ"], [Cm(10.3)] + [Cm(10.3), Cm(12.3)])
field_pair(doc, "๗) เพื่อนสนิท หรือผู้ให้พักอาศัย", "{suspect_friends}")
field_pair(doc, "๘) ที่อยู่ปัจจุบัน", "{suspect_current_address}", "ความสูง", "{phys_height}")
field_pair(doc, "๙) สถานที่ชอบไปเที่ยวเตร่", "{suspect_hangouts}", "น้ำหนัก", "{phys_weight}")
field_pair(doc, "๑๐) อาวุธที่ใช้ในการกระทำผิด", "{crime_weapon}", "สีผิว", "{phys_skin}")
field_pair(doc, "๑๑) ยานพาหนะที่ใช้", "{crime_vehicle}", "ศีรษะ", "{phys_head}")
field_pair(doc, "๑๒) ผู้ร่วมกระทำผิด", "{crime_accomplice}", "ผม", "{phys_hair}")
field_pair(doc, "๑๓) รายละเอียดวิธีกระทำความผิด", "{crime_details}", "รูปหน้า", "{phys_face}")
field_pair(doc, "๑๔) พฤติการณ์ที่น่าสนใจ", "{crime_notes}", "คิ้ว,ตา", "{phys_eyes}")
field_pair(doc, "ตำรวจหรือบุคคลผู้รู้จักตัว", "{known_by}", "จมูก", "{phys_nose}")
field_pair(doc, "บัตรประจำตัวเลขที่", "{suspect_id_number}", "ใบหู", "{phys_ears}")
field_pair(doc, "ออกให้อำเภอ", "{id_issue_amphoe}", "หนวดเครา", "{phys_beard}")
field_pair(doc, "จังหวัด", "{id_issue_province}", "ตำหนิ", "{phys_marks}")
field_pair(doc, "๑๕) ประวัติการกระทำผิด", "{crime_history}", "สำเนียง", "{phys_accent}")

# page break, then the 4-angle photo page
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.0
p.add_run().add_break(WD_BREAK.PAGE)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("ภาพประกอบแนบท้ายผู้ถูกจับและควบคุม ตามข้อ (๑)"); set_font(r, bold=True)
add_p(doc, "")
photo_table = doc.add_table(rows=2, cols=2)
photo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
photo_order = [("front", os.path.join(PLACEHOLDER_DIR, "front.jpg")), ("back", os.path.join(PLACEHOLDER_DIR, "back.jpg")),
               ("left", os.path.join(PLACEHOLDER_DIR, "left.jpg")), ("right", os.path.join(PLACEHOLDER_DIR, "right.jpg"))]
cells_flat = [photo_table.rows[0].cells[0], photo_table.rows[0].cells[1],
              photo_table.rows[1].cells[0], photo_table.rows[1].cells[1]]
for (label, path), tcell in zip(photo_order, cells_flat):
    para = tcell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=CmSize(7.6), height=CmSize(9.2))
add_p(doc, "")
add_p(doc, "{suspect_title}{suspect_full_name} หรือ {suspect_alias} อายุ {suspect_age} ปี เลขบัตรประจำตัวประชาชน {suspect_id_number} "
          "อยู่บ้านเลขที่ {suspect_address} {charge}", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_p(doc, "")
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run("( ............................................ )"); set_font(r)
p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p4.add_run("{suspect_full_name}"); set_font(r)
doc.save(OUT + "profile_and_photos_template.docx")

# ---------- 3. court referral ----------
doc = base_doc()
letterhead(doc, "{doc_number}")
add_mixed(doc, [("เรื่อง\t", True), ("ส่งตัวผู้ต้องหาตามหมายจับของ{court_name}", False)])
add_mixed(doc, [("เรียน\t", True), ("ผู้พิพากษาหัวหน้า{court_name}", False)])
add_p(doc, "อ้างถึง  จับตามหมายจับของ{court_name} ที่ {warrant_no} คดีหมายเลขดำที่ {case_black_no} ลงวันที่ {warrant_date} ในความผิดฐาน \"{charge}\"")
add_p(doc, "สิ่งที่ส่งมาด้วย  บันทึกจับกุมตัว จำนวน ๑ ฉบับ")
add_p(doc, "ด้วยเมื่อวันที่ {arrest_date} {court_name} ที่ {warrant_no} คดีหมายเลขดำที่ {case_black_no} ลงวันที่ {warrant_date} ในความผิดฐาน {charge} นั้น")
add_p(doc, "{station_name} จังหวัด{province} ขอเรียนว่า ได้ทำการจับกุม {suspect_title}{suspect_full_name} อายุ {suspect_age} ปี "
          "สัญชาติ {suspect_nationality} หมายเลขประจำตัวประชาชน {suspect_id_number} ที่อยู่ {suspect_address} "
          "จึงมอบหมายให้ {officer1_rank}{officer1_name} {officer1_position} {station_name} จังหวัด{province} พร้อมพวก "
          "เป็นผู้นำตัว {suspect_title}{suspect_full_name} มาส่งตัวที่{court_name} เพื่อดำเนินการต่อไป "
          "พร้อมนี้ ได้แนบบันทึกการจับกุมตัว มาพร้อมนี้ด้วยแล้ว จำนวน ๑ ฉบับ")
add_p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา")
add_p(doc, "ขอแสดงความนับถือ")
signature_block(doc, "{signer_rank}", "{signer_name}", "{signer_position}")
add_p(doc, "{station_name} จังหวัด{province}")
doc.save(OUT + "court_referral_template.docx")

# ---------- 4. phone consent (m.80) ----------
doc = base_doc()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("บันทึกความยินยอมให้ถอดข้อมูลในโทรศัพท์มือถือและข้อมูลอิเล็กทรอนิกส์"); set_font(r, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ตามมาตรา ๘๐ พ.ร.บ.คุ้มครองข้อมูลส่วนบุคคล พ.ศ.๒๕๖๒"); set_font(r, bold=True)
add_p(doc, "")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("สถานที่บันทึก {record_place}"); set_font(r)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("วันที่ {record_date}"); set_font(r)
add_p(doc, "บันทึกฉบับนี้จัดทำไว้เพื่อแสดงว่า ข้าพเจ้า {suspect_title}{suspect_full_name} อายุ {suspect_age} ปี "
          "ที่อยู่ {suspect_address} เลขประจำตัวประชาชน {suspect_id_number}")
add_p(doc, "เป็นเจ้าของ/ผู้ครอบครอง/ผู้ใช้ โทรศัพท์เคลื่อนที่ ยี่ห้อ {phone_brand} รุ่น {phone_model} สี {phone_color}")
field_line(doc, "เลขอีมี่ (1)", "{phone_imei1}   เลขอีมี่ (2) {phone_imei2}")
field_line(doc, "ซิมการ์ดหมายเลขโทร (1)", "{sim1_number} เครือข่าย {sim1_network}")
field_line(doc, "ซิมการ์ดหมายเลขโทร (2)", "{sim2_number} เครือข่าย {sim2_network}")
add_p(doc, "ยินยอมให้เจ้าหน้าที่ตำรวจถอดข้อมูลในโทรศัพท์เคลื่อนที่และข้อมูลทางอิเล็กทรอนิกส์ เช่น หมายเลขโทรศัพท์, "
          "รายชื่อ, ข้อมูลการใช้, ข้อความ, ภาพถ่าย, ข้อมูลแอปพลิเคชันไลน์, ข้อมูลแอปพลิเคชันเฟซบุ๊ก, ธุรกรรมการเงินทางอิเล็กทรอนิกส์, อีเมล")
add_p(doc, "อนึ่ง ก่อนถอดข้อมูล เจ้าหน้าที่ตำรวจให้ทราบแล้วว่าจะนำข้อมูลที่ได้ใช้เป็นพยานหลักฐานประกอบการสืบสวนสอบสวน "
          "และข้าพเจ้าให้ความยินยอมโดยไม่ได้ถูกบังคับ ข่มขู่ ขู่เข็ญ หลอกลวง หรือกระทำให้เกิดความเกรงกลัวว่าจะเกิดอันตรายแก่ชีวิต ร่างกาย "
          "ทรัพย์สิน ชื่อเสียงของตนเองหรือบุคคลในครอบครัวแต่ประการใด")
add_p(doc, "เจ้าหน้าที่ตำรวจได้อ่านบันทึกให้ฟังแล้ว รับว่าถูกต้องจึงลงลายมือชื่อไว้เป็นหลักฐาน")
add_p(doc, "")
for label in ["(ลงชื่อ) ............................................. ผู้ยินยอม", "(ลงชื่อ) ............................................. พยาน",
              "(ลงชื่อ) {officer2_rank}{officer2_name} ............................................. บันทึก/อ่าน"]:
    p = doc.add_paragraph(); r = p.add_run(label); set_font(r)
doc.save(OUT + "phone_consent_m80_template.docx")

# ---------- 5. arrest report (new, general letter shell) ----------
doc = base_doc()
letterhead(doc, "{doc_number}")
add_mixed(doc, [("เรื่อง\t", True), ("บันทึกการจับกุม", False)])
add_mixed(doc, [("เรียน\t", True), ("พนักงานสอบสวน{station_name}", False)])
add_mixed(doc, [("อ้างถึง\t", True), ("{warrant_reference}", False)])
p = doc.add_paragraph(); r = p.add_run("ภาคเหตุ"); set_font(r, bold=True)
add_p(doc, "ด้วยเมื่อวันที่ {arrest_date} เวลา {arrest_time} น. เจ้าหน้าที่ตำรวจ {station_name} ประกอบด้วย "
          "{officer1_rank}{officer1_name} {officer1_position}, {officer2_rank}{officer2_name} {officer2_position} และ "
          "{officer3_rank}{officer3_name} {officer3_position} ได้ร่วมกันจับกุม {suspect_title}{suspect_full_name} อายุ {suspect_age} ปี "
          "สัญชาติ {suspect_nationality} อยู่บ้านเลขที่ {suspect_address} ณ {arrest_location} "
          "ในข้อหา \"{charge}\" {warrant_reference}")
p = doc.add_paragraph(); r = p.add_run("ภาคความประสงค์"); set_font(r, bold=True)
add_p(doc, "เจ้าหน้าที่ผู้จับกุมจึงนำตัว {suspect_title}{suspect_full_name} พร้อมของกลาง (ถ้ามี) ส่งพนักงานสอบสวน{station_name} "
          "เพื่อดำเนินคดีตามกฎหมายต่อไป")
p = doc.add_paragraph(); r = p.add_run("ภาคสรุป"); set_font(r, bold=True)
add_p(doc, "จึงบันทึกไว้เป็นหลักฐาน และให้ผู้ถูกจับกุมลงลายมือชื่อรับทราบข้อกล่าวหาไว้ ณ ที่นี้")
add_p(doc, "")
for tag in ["officer1", "officer2", "officer3"]:
    p = doc.add_paragraph()
    r = p.add_run("(ลงชื่อ) {%s_rank}{%s_name} ผู้จับกุม" % (tag, tag))
    set_font(r)
p = doc.add_paragraph()
r = p.add_run("(ลงชื่อ) ............................................. ผู้ถูกจับกุม")
set_font(r)
doc.save(OUT + "arrest_report_template.docx")

print("done")
